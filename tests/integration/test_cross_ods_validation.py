"""Pruebas de la barrera anti "ODS equivocada" en el flujo completo.

Incluye casos sintéticos (corren también en CI) y el cruce con los archivos
reales: el Excel de la ODS 266 (HOCOL) contra la plantilla Word de la ODS 11.
"""

from __future__ import annotations

from pathlib import Path

import docx
import openpyxl
import pytest

from ods_reporter.application.use_cases.process_ods import ProcessRequest
from ods_reporter.infrastructure.matching.roman_numerals import int_to_roman
from tests.integration.test_process_ods import FakeProgress, _make_use_case

_WORD_LABELS = [
    "Identificar y analizar requisitos ambientales",
    "Elaborar el recurso de reposición cuando aplique",
    "Documentar conceptos técnicos ambientales",
    "Acompañar visitas de campo y comités",
]

_OTHER_LABELS = [
    "Informe de análisis de variables del sistema",
    "Matriz de control documental",
    "Reporte de auditoría y plan de acción",
    "Gestión de peticiones hasta el cierre",
]


def _build_word(path: Path, labels: list[str]) -> None:
    document = docx.Document()
    table = document.add_table(rows=1 + len(labels), cols=2)
    table.rows[0].cells[0].text = "No"
    table.rows[0].cells[1].text = "Actividades"
    for index, label in enumerate(labels, start=1):
        cells = table.rows[index].cells
        cells[0].text = int_to_roman(index)
        cell = cells[1]
        cell.paragraphs[0].text = f"Actividad: {label}"
        cell.add_paragraph(f"Descripción del entregable: Informe de {label.lower()}")
        cell.add_paragraph("Descripción de las actividades realizadas:")
        cell.add_paragraph("", style="List Bullet")
    document.save(str(path))


def _build_excel(path: Path, labels: list[str]) -> None:
    workbook = openpyxl.Workbook()
    sheet = workbook.active
    sheet.title = "MARZO"
    sheet["A2"] = "Profesional responsable:"
    sheet["B2"] = "Ana Pérez"
    sheet["A5"] = "ID"
    sheet["B5"] = "ACTIVIDADES"
    sheet["C5"] = "ENTREGABLES"
    sheet["D5"] = "DESCRIPCIÓN DE ACTIVIDADES O ENTREGABLES REALIZADOS EN EL PERIODO"
    for index, label in enumerate(labels, start=1):
        row = 5 + index
        sheet.cell(row=row, column=1, value=index)
        sheet.cell(row=row, column=2, value=label)
        sheet.cell(row=row, column=3, value=f"Informe de {label.lower()}")
        sheet.cell(row=row, column=4, value="Se realizó la actividad")
    workbook.save(str(path))


def _request(word: Path, excel: Path, tmp_path: Path, name: str) -> ProcessRequest:
    return ProcessRequest(
        word_template=word,
        excel_files=(excel,),
        output_dir=tmp_path,
        month="MARZO",
        output_name=name,
    )


def test_synthetic_cross_ods_file_is_rejected(tmp_path: Path) -> None:
    word = tmp_path / "plantilla.docx"
    excel = tmp_path / "otro_proyecto.xlsx"
    _build_word(word, _WORD_LABELS)
    _build_excel(excel, _OTHER_LABELS)

    use_case = _make_use_case(FakeProgress())
    plan = use_case.plan(_request(word, excel, tmp_path, "cruce.docx")).unwrap()

    assert plan.professionals == ()
    assert plan.planned == ()
    assert any("no corresponde a esta ODS" in e for e in plan.read_errors)


def test_synthetic_same_ods_is_accepted(tmp_path: Path) -> None:
    word = tmp_path / "plantilla.docx"
    excel = tmp_path / "mismo_proyecto.xlsx"
    _build_word(word, _WORD_LABELS)
    _build_excel(excel, _WORD_LABELS)

    use_case = _make_use_case(FakeProgress())
    plan = use_case.plan(_request(word, excel, tmp_path, "mismo.docx")).unwrap()

    assert plan.read_errors == ()
    assert len(plan.professionals) == 1
    assert len(plan.planned) == len(_WORD_LABELS)
    # Enunciados idénticos: sin advertencias de "no coincide".
    assert not any("no coincide" in w for w in plan.warnings)


def test_real_cross_ods_excel_is_rejected(word_blank_fixture: Path, tmp_path: Path) -> None:
    ods266 = Path("tests/fixtures/ODS266_Excel.xlsx")
    if not ods266.exists():
        pytest.skip("Falta el fixture real de la ODS 266")

    use_case = _make_use_case(FakeProgress())
    plan = use_case.plan(
        _request(word_blank_fixture, ods266, tmp_path, "cruce_real.docx")
    ).unwrap()

    assert plan.professionals == ()
    assert any("no corresponde a esta ODS" in e for e in plan.read_errors)
