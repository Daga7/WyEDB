"""Prueba end-to-end del caso de uso ProcessODSUseCase con archivos reales."""

from __future__ import annotations

from pathlib import Path

import docx
import pytest

from ods_reporter.application.ports.progress_port import EventLevel
from ods_reporter.application.use_cases.process_ods import ProcessODSUseCase, ProcessRequest
from ods_reporter.infrastructure.excel.openpyxl_reader import OpenpyxlExcelReader
from ods_reporter.infrastructure.filesystem.file_service import FileService
from ods_reporter.infrastructure.reporting.report_writer import ReportWriter
from ods_reporter.infrastructure.word.docx_processor import DocxProcessor
from ods_reporter.shared.constants import DEFAULT_EMPTY_ACTIVITY_TEXT


class FakeProgress:
    """Implementación de ProgressPort para pruebas (registra eventos en memoria)."""

    def __init__(self, cancel_after: int | None = None) -> None:
        self.events: list[tuple[EventLevel, str]] = []
        self.last_progress: tuple[int, int] | None = None
        self._cancel_after = cancel_after
        self._progress_calls = 0

    def event(self, level: EventLevel, message: str) -> None:
        self.events.append((level, message))

    def progress(self, current: int, total: int) -> None:
        self.last_progress = (current, total)
        self._progress_calls += 1

    def is_cancelled(self) -> bool:
        if self._cancel_after is None:
            return False
        return self._progress_calls >= self._cancel_after


def _make_use_case(progress: FakeProgress) -> ProcessODSUseCase:
    return ProcessODSUseCase(
        excel_reader=OpenpyxlExcelReader(),
        word_processor=DocxProcessor(),
        file_service=FileService(),
        progress=progress,
        report_writer=ReportWriter(),
    )


def test_end_to_end_fills_word(
    excel_fixture: Path, word_fixture: Path, tmp_path: Path
) -> None:
    progress = FakeProgress()
    use_case = _make_use_case(progress)
    request = ProcessRequest(
        word_template=word_fixture,
        excel_files=(excel_fixture,),
        output_dir=tmp_path,
        month="MAYO",
        output_name="resultado.docx",
    )

    result = use_case.execute(request)

    assert result.is_ok()
    processing = result.unwrap()
    assert processing.professionals_processed == 1
    # En mayo hay 9 actividades con contenido.
    assert processing.activities_with_content == 9
    assert processing.items_written > 0
    assert processing.activities_not_found == 0

    # El original NO se modificó; la salida sí existe.
    output = tmp_path / "resultado.docx"
    assert output.exists()
    assert word_fixture.read_bytes() != output.read_bytes()

    # Se generó el reporte de resumen junto a la salida.
    report = tmp_path / "resultado_reporte.txt"
    assert report.exists()
    assert "Resumen del procesamiento" in report.read_text(encoding="utf-8")
    assert processing.summary != ""


def test_output_contains_inserted_content(
    excel_fixture: Path, word_blank_fixture: Path, tmp_path: Path
) -> None:
    progress = FakeProgress()
    use_case = _make_use_case(progress)
    request = ProcessRequest(
        word_template=word_blank_fixture,
        excel_files=(excel_fixture,),
        output_dir=tmp_path,
        month="MAYO",
        output_name="resultado.docx",
    )
    use_case.execute(request)

    document = docx.Document(str(tmp_path / "resultado.docx"))
    full_text = "\n".join(
        p.text for t in document.tables for row in t.rows for c in row.cells for p in c.paragraphs
    )
    assert "Apoyo en Radicación de 3 ICA" in full_text


def test_multiple_excels_accumulate_content(
    excel_fixture: Path, word_blank_fixture: Path, tmp_path: Path
) -> None:
    # Caso real: varios profesionales (varios Excel) en una sola ejecución.
    # Se usa el mismo Excel 3 veces para simular 3 profesionales.
    progress = FakeProgress()
    use_case = _make_use_case(progress)
    request = ProcessRequest(
        word_template=word_blank_fixture,
        excel_files=(excel_fixture, excel_fixture, excel_fixture),
        output_dir=tmp_path,
        month="MAYO",
        output_name="multi.docx",
    )
    result = use_case.execute(request).unwrap()

    assert result.professionals_processed == 3
    assert result.activities_with_content == 27  # 9 por profesional

    document = docx.Document(str(tmp_path / "multi.docx"))
    full_text = "\n".join(
        p.text for t in document.tables for row in t.rows for c in row.cells for p in c.paragraphs
    )
    # El contenido de cada profesional se acumula (uno debajo del otro).
    assert full_text.count("Apoyo en Radicación de 3 ICA") == 3


def test_empty_month_fills_all_with_default(
    excel_fixture: Path, word_fixture: Path, tmp_path: Path
) -> None:
    # JUNIO no tiene contenido: todos los slots deben recibir el texto por defecto.
    progress = FakeProgress()
    use_case = _make_use_case(progress)
    request = ProcessRequest(
        word_template=word_fixture,
        excel_files=(excel_fixture,),
        output_dir=tmp_path,
        month="JUNIO",
        output_name="junio.docx",
    )
    result = use_case.execute(request)
    processing = result.unwrap()
    assert processing.activities_with_content == 0
    assert processing.default_slots_filled > 0

    document = docx.Document(str(tmp_path / "junio.docx"))
    full_text = "\n".join(
        p.text for t in document.tables for row in t.rows for p in row.cells[1].paragraphs
        if len(row.cells) > 1
    )
    assert DEFAULT_EMPTY_ACTIVITY_TEXT in full_text


def test_end_to_end_with_blank_template(
    excel_fixture: Path, word_blank_fixture: Path, tmp_path: Path
) -> None:
    # Verificación realista: plantilla EN BLANCO + Excel real de mayo.
    progress = FakeProgress()
    use_case = _make_use_case(progress)
    request = ProcessRequest(
        word_template=word_blank_fixture,
        excel_files=(excel_fixture,),
        output_dir=tmp_path,
        month="MAYO",
        output_name="blank_resultado.docx",
    )
    result = use_case.execute(request).unwrap()

    assert result.activities_with_content == 9
    assert result.activities_not_found == 0
    assert result.items_written >= 12
    # En una plantilla limpia, los slots vacíos reciben el texto por defecto.
    assert result.default_slots_filled > 0

    document = docx.Document(str(tmp_path / "blank_resultado.docx"))
    full_text = "\n".join(
        p.text for t in document.tables for row in t.rows for c in row.cells for p in c.paragraphs
    )
    assert "Apoyo en Radicación de 3 ICA" in full_text
    assert DEFAULT_EMPTY_ACTIVITY_TEXT in full_text


def test_single_professional_multiple_entregables(
    excel_fixture: Path, word_blank_fixture: Path, tmp_path: Path
) -> None:
    # Requisito: una actividad con VARIOS entregables de UN solo profesional debe
    # colocar el contenido de cada entregable en su sub-fila correspondiente.
    import docx as _docx

    from ods_reporter.infrastructure.word.docx_reader import DocxReader
    from ods_reporter.shared.text_utils import normalize_text

    progress = FakeProgress()
    use_case = _make_use_case(progress)
    use_case.execute(
        ProcessRequest(
            word_template=word_blank_fixture,
            excel_files=(excel_fixture,),
            output_dir=tmp_path,
            month="MAYO",
            output_name="multi_ent.docx",
        )
    )

    document = _docx.Document(str(tmp_path / "multi_ent.docx"))
    act8 = next(a for a in DocxReader().read_activities(document) if a.ordinal == 8)
    assert len(act8.entregables) == 2

    # Cada entregable tiene su propio contenido (distinto), no mezclado.
    contents = []
    for entregable in act8.entregables:
        cell = entregable.cell
        idx = next(
            (
                k
                for k, p in enumerate(cell.paragraphs)
                if "actividades realizadas" in normalize_text(p.text)
            ),
            -1,
        )
        items = [p.text.strip() for p in cell.paragraphs[idx + 1 :] if p.text.strip()]
        contents.append(" ".join(items))

    assert any("socializa presentación" in c.lower() for c in contents)
    assert any("socialización de medidas" in c.lower() for c in contents)
    # No están mezclados en el mismo entregable.
    assert contents[0] != contents[1]


def test_bad_excel_reports_file_and_continues(
    excel_fixture: Path, word_blank_fixture: Path, tmp_path: Path
) -> None:
    # Robustez: un Excel corrupto no detiene el resto; el error nombra el archivo.
    bad = tmp_path / "corrupto.xlsx"
    bad.write_bytes(b"esto no es un xlsx valido")

    progress = FakeProgress()
    use_case = _make_use_case(progress)
    result = use_case.execute(
        ProcessRequest(
            word_template=word_blank_fixture,
            excel_files=(bad, excel_fixture),  # uno malo + uno bueno
            output_dir=tmp_path,
            month="MAYO",
            output_name="mixto.docx",
        )
    ).unwrap()

    # El archivo bueno se procesó pese al malo.
    assert result.professionals_processed == 1
    assert result.activities_with_content > 0
    # Hay un error y menciona el archivo culpable.
    assert result.has_errors
    assert any("corrupto.xlsx" in e for e in result.errors)
    # El resumen incluye la guía y el nombre del archivo.
    assert "corrupto.xlsx" in result.summary


def test_invalid_input_returns_err(word_fixture: Path, tmp_path: Path) -> None:
    progress = FakeProgress()
    use_case = _make_use_case(progress)
    request = ProcessRequest(
        word_template=word_fixture,
        excel_files=(),  # sin Excel
        output_dir=tmp_path,
        month="MAYO",
    )
    result = use_case.execute(request)
    assert result.is_err()


# --- Vista previa: fases plan y apply ---

def _plan_request(
    excel: Path, word: Path, tmp_path: Path, output_name: str
) -> ProcessRequest:
    return ProcessRequest(
        word_template=word,
        excel_files=(excel,),
        output_dir=tmp_path,
        month="MAYO",
        output_name=output_name,
    )


def test_plan_reports_structure_without_writing(
    excel_fixture: Path, word_blank_fixture: Path, tmp_path: Path
) -> None:
    use_case = _make_use_case(FakeProgress())
    request = _plan_request(excel_fixture, word_blank_fixture, tmp_path, "plan.docx")

    plan = use_case.plan(request).unwrap()

    assert len(plan.professionals) == 1
    assert len(plan.planned) == 9  # las 9 actividades con contenido de mayo
    assert all(p.matched for p in plan.planned)
    assert plan.unmatched == ()
    assert len(plan.word_activities) > 0
    assert plan.read_errors == ()
    # La plantilla ODS11 tiene la sección de observaciones, pero mayo no trae
    # actividades adicionales diligenciadas.
    assert plan.word_has_other_section is True
    assert plan.other_activities_count == 0
    # La fase de plan NO genera el documento.
    assert not (tmp_path / "plan.docx").exists()


def test_apply_with_omit_override_skips_activity(
    excel_fixture: Path, word_blank_fixture: Path, tmp_path: Path
) -> None:
    use_case = _make_use_case(FakeProgress())
    request = _plan_request(excel_fixture, word_blank_fixture, tmp_path, "omitida.docx")
    plan = use_case.plan(request).unwrap()
    first = plan.planned[0]

    result = use_case.apply(request, plan, {first.key: None}).unwrap()

    assert result.activities_with_content == 8  # una fue omitida por el usuario
    assert any("omitida por el usuario" in w for w in result.warnings)
    assert (tmp_path / "omitida.docx").exists()


def test_apply_with_redirect_override_moves_content(
    excel_fixture: Path, word_blank_fixture: Path, tmp_path: Path
) -> None:
    from ods_reporter.infrastructure.word.docx_reader import DocxReader

    use_case = _make_use_case(FakeProgress())
    request = _plan_request(excel_fixture, word_blank_fixture, tmp_path, "movida.docx")
    plan = use_case.plan(request).unwrap()

    first = plan.planned[0]
    professional = plan.professionals[first.professional_index]
    activity = next(a for a in professional.activities if a.ordinal == first.ordinal)
    item_text = activity.all_content_items[0].text
    target = next(
        o.ordinal for o in plan.word_activities if o.ordinal != first.ordinal
    )

    result = use_case.apply(request, plan, {first.key: target}).unwrap()

    assert any("manualmente" in w for w in result.warnings)

    document = docx.Document(str(tmp_path / "movida.docx"))
    by_ordinal = {a.ordinal: a for a in DocxReader().read_activities(document)}

    def cell_text(ordinal: int) -> str:
        return "\n".join(
            p.text
            for entregable in by_ordinal[ordinal].entregables
            for p in entregable.cell.paragraphs
        )

    assert item_text in cell_text(target)
    assert item_text not in cell_text(first.ordinal)
