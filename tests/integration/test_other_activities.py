"""Pruebas del flujo de "Otras actividades solicitadas" (Excel -> Word).

Usa un Excel y un Word sintéticos construidos al vuelo, de modo que corren
también donde no están los fixtures reales (p. ej. en CI).
"""

from __future__ import annotations

from datetime import datetime
from pathlib import Path

import docx
import openpyxl
import pytest

from ods_reporter.application.services.report_builder import ReportBuilder
from ods_reporter.application.use_cases.process_ods import ProcessODSUseCase, ProcessRequest
from ods_reporter.infrastructure.excel.openpyxl_reader import OpenpyxlExcelReader
from ods_reporter.infrastructure.filesystem.file_service import FileService
from ods_reporter.infrastructure.word.docx_processor import DocxProcessor
from tests.integration.test_docx_processor_plan import _build_word
from tests.integration.test_process_ods import FakeProgress


def _build_excel(path: Path) -> None:
    """Excel mínimo: una actividad principal + sección de otras actividades."""
    workbook = openpyxl.Workbook()
    sheet = workbook.active
    sheet.title = "MARZO"

    sheet["A2"] = "Profesional responsable:"
    sheet["B2"] = "Ana Pérez"

    sheet["A5"] = "ID"
    sheet["B5"] = "ACTIVIDADES"
    sheet["C5"] = "ENTREGABLES"
    sheet["D5"] = "DESCRIPCIÓN DE ACTIVIDADES O ENTREGABLES REALIZADOS EN EL PERIODO"

    sheet["A6"] = 1
    sheet["B6"] = "Asesoría ambiental en campo"
    sheet["C6"] = "Informe de asesoría ambiental en campo"
    sheet["D6"] = "7_Se asesoró la obra A\n8_Se revisó el plan B"

    # La fila del marcador es también el encabezado de la sección.
    sheet["A8"] = "ID"
    sheet["B8"] = "OTRAS ACTIVIDADES SOLICITADAS POR ECOPETROL"
    sheet["C8"] = "FECHA DE EJECUCIÓN"

    sheet["A9"] = 1
    sheet["B9"] = "Capacitación extra al equipo"
    sheet["C9"] = datetime(2026, 3, 16)

    sheet["A10"] = 2
    sheet["B10"] = "No se requirió esta actividad"

    sheet["A11"] = 3  # fila con ID pero sin descripción: se ignora

    sheet["A12"] = "FIN DEL REPORTE"

    workbook.save(str(path))


@pytest.fixture
def excel_path(tmp_path: Path) -> Path:
    path = tmp_path / "reporte.xlsx"
    _build_excel(path)
    return path


def test_reader_extracts_other_activities(excel_path: Path) -> None:
    raw = OpenpyxlExcelReader().read_month(excel_path, "MARZO")

    # Extracción fiel: incluye la fila de "no se requirió" (el filtro es del builder).
    assert [(o.text, o.date) for o in raw.other_activities] == [
        ("Capacitación extra al equipo", "16/03/2026"),
        ("No se requirió esta actividad", ""),
    ]


def test_builder_filters_and_adds_date(excel_path: Path) -> None:
    raw = OpenpyxlExcelReader().read_month(excel_path, "MARZO")
    professional = ReportBuilder().build(raw)

    assert [i.text for i in professional.other_activities] == [
        "Capacitación extra al equipo (16/03/2026)"
    ]
    # Y los numerales "7_" / "8_" del contenido regular quedaron limpios.
    activity = professional.activities[0]
    assert [i.text for i in activity.all_content_items] == [
        "Se asesoró la obra A",
        "Se revisó el plan B",
    ]


def test_end_to_end_inserts_other_activities_in_word(
    excel_path: Path, tmp_path: Path
) -> None:
    word_path = tmp_path / "plantilla.docx"
    _build_word(word_path, with_observaciones=True)

    use_case = ProcessODSUseCase(
        excel_reader=OpenpyxlExcelReader(),
        word_processor=DocxProcessor(),
        file_service=FileService(),
        progress=FakeProgress(),
    )
    request = ProcessRequest(
        word_template=word_path,
        excel_files=(excel_path,),
        output_dir=tmp_path,
        month="MARZO",
        output_name="salida.docx",
    )

    plan = use_case.plan(request).unwrap()
    assert plan.other_activities_count == 1
    assert plan.word_has_other_section is True
    assert plan.general_warnings == ()

    result = use_case.apply(request, plan, {}).unwrap()
    assert result.other_activities_written == 1
    assert "Actividades adicionales  : 1" in result.summary

    document = docx.Document(str(tmp_path / "salida.docx"))
    full_text = "\n".join(
        p.text for t in document.tables for row in t.rows for c in row.cells for p in c.paragraphs
    )
    assert "Capacitación extra al equipo (16/03/2026)" in full_text
    assert "No se requirió esta actividad" not in full_text


def test_plan_warns_when_word_has_no_section(excel_path: Path, tmp_path: Path) -> None:
    word_path = tmp_path / "plantilla_sin_obs.docx"
    _build_word(word_path, with_observaciones=False)

    use_case = ProcessODSUseCase(
        excel_reader=OpenpyxlExcelReader(),
        word_processor=DocxProcessor(),
        file_service=FileService(),
        progress=FakeProgress(),
    )
    request = ProcessRequest(
        word_template=word_path,
        excel_files=(excel_path,),
        output_dir=tmp_path,
        month="MARZO",
        output_name="salida_sin_obs.docx",
    )

    plan = use_case.plan(request).unwrap()
    assert plan.word_has_other_section is False
    assert any("se omitirán" in w for w in plan.general_warnings)

    result = use_case.apply(request, plan, {}).unwrap()
    assert result.other_activities_written == 0
    assert any("se omitieron" in w for w in result.warnings)
