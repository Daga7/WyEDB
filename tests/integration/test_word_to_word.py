"""Pruebas del modo Word → Word: leer el Word de un profesional y volcarlo
en la plantilla emparejando por ENUNCIADO (no solo por numeral).

Los documentos se construyen sintéticamente para que las pruebas corran
también sin los fixtures reales (p. ej. en CI).
"""

from __future__ import annotations

from pathlib import Path

import docx

from ods_reporter.application.use_cases.process_ods import (
    ProcessODSUseCase,
    ProcessRequest,
)
from ods_reporter.infrastructure.filesystem.file_service import FileService
from ods_reporter.infrastructure.matching.roman_numerals import int_to_roman
from ods_reporter.infrastructure.readers.report_reader_router import ReportReaderRouter
from ods_reporter.infrastructure.word.docx_processor import DocxProcessor
from ods_reporter.infrastructure.word.docx_reader import DocxReader
from ods_reporter.infrastructure.word.docx_source_reader import DocxSourceReader
from ods_reporter.shared.constants import DEFAULT_EMPTY_ACTIVITY_TEXT
from tests.integration.test_process_ods import FakeProgress

_LABELS = [
    "Identificar y analizar requisitos ambientales",
    "Elaborar el recurso de reposición cuando aplique",
    "Documentar conceptos técnicos ambientales",
    "Acompañar visitas de campo y comités",
]


def _build_template(path: Path, labels: list[str]) -> None:
    """Plantilla en blanco: slots vacíos, numerales romanos en orden."""
    document = docx.Document()
    table = document.add_table(rows=1 + len(labels), cols=2)
    table.rows[0].cells[0].text = "No"
    table.rows[0].cells[1].text = "Actividades"
    for index, label in enumerate(labels, start=1):
        cells = table.rows[index].cells
        cells[0].text = int_to_roman(index)
        cell = cells[1]
        cell.paragraphs[0].text = f"Actividad: {label}"
        cell.add_paragraph(f"Descripción del entregable: Informe de {label.lower()}")
        cell.add_paragraph("Descripción de las actividades realizadas:")
        cell.add_paragraph("", style="List Bullet")
    document.save(str(path))


def _build_source(
    path: Path,
    entries: list[tuple[str, list[str]]],
    *,
    professional: str = "Ana Torres",
) -> None:
    """Word diligenciado por un profesional: viñetas con contenido."""
    document = docx.Document()
    document.add_paragraph(f"Profesional responsable: {professional}")
    table = document.add_table(rows=1 + len(entries), cols=2)
    table.rows[0].cells[0].text = "No"
    table.rows[0].cells[1].text = "Actividades"
    for index, (label, items) in enumerate(entries, start=1):
        cells = table.rows[index].cells
        cells[0].text = int_to_roman(index)
        cell = cells[1]
        cell.paragraphs[0].text = f"Actividad: {label}"
        cell.add_paragraph(f"Descripción del entregable: Informe de {label.lower()}")
        cell.add_paragraph("Descripción de las actividades realizadas:")
        if not items:
            cell.add_paragraph("", style="List Bullet")
        for item in items:
            cell.add_paragraph(item, style="List Bullet")
    document.save(str(path))


def _make_use_case(progress: FakeProgress) -> ProcessODSUseCase:
    return ProcessODSUseCase(
        excel_reader=ReportReaderRouter(),
        word_processor=DocxProcessor(),
        file_service=FileService(),
        progress=progress,
    )


def _request(word: Path, source: Path, tmp_path: Path, name: str) -> ProcessRequest:
    return ProcessRequest(
        word_template=word,
        excel_files=(source,),
        output_dir=tmp_path,
        month="MARZO",
        output_name=name,
    )


def _activity_texts(path: Path) -> dict[int, str]:
    document = docx.Document(str(path))
    activities = DocxReader().read_activities(document)
    return {
        a.ordinal: "\n".join(
            p.text for e in a.entregables for p in e.cell.paragraphs
        )
        for a in activities
    }


# --- Lector de Word de profesional ---


def test_source_reader_extracts_report(tmp_path: Path) -> None:
    source = tmp_path / "reporte_ana.docx"
    _build_source(
        source,
        [(_LABELS[0], ["Se asesoró la obra A", "Se revisó el permiso B"])],
    )

    raw = DocxSourceReader().read_month(source, "MARZO")

    assert raw.professional_name == "Ana Torres"
    assert raw.source_file == "reporte_ana.docx"
    assert len(raw.activities) == 1
    content = raw.activities[0].entregables[0].raw_content
    assert "Se asesoró la obra A" in content
    assert "Se revisó el permiso B" in content


def test_source_reader_falls_back_to_file_name(tmp_path: Path) -> None:
    source = tmp_path / "Informe Carlos Pérez.docx"
    document = docx.Document()
    table = document.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "No"
    table.rows[0].cells[1].text = "Actividades"
    cells = table.rows[1].cells
    cells[0].text = "I"
    cell = cells[1]
    cell.paragraphs[0].text = f"Actividad: {_LABELS[0]}"
    cell.add_paragraph("Algo realizado", style="List Bullet")
    document.save(str(source))

    raw = DocxSourceReader().read_month(source, "MARZO")
    assert raw.professional_name == "Informe Carlos Pérez"


# --- Enrutador ---


def test_router_dispatches_by_extension(tmp_path: Path) -> None:
    calls: list[str] = []

    class _Fake:
        def __init__(self, tag: str) -> None:
            self._tag = tag

        def read_month(self, file_path: Path, month: str):  # noqa: ANN201
            calls.append(self._tag)
            return None

    router = ReportReaderRouter(excel_reader=_Fake("excel"), word_reader=_Fake("word"))
    router.read_month(Path("a.docx"), "MARZO")
    router.read_month(Path("b.xlsx"), "MARZO")
    router.read_month(Path("c.XLSM"), "MARZO")
    assert calls == ["word", "excel", "excel"]


# --- Flujo completo: emparejado por enunciado ---


def test_word_source_matches_by_label_even_with_shuffled_numerals(
    tmp_path: Path,
) -> None:
    """El profesional numeró distinto: el contenido cae donde dice el ENUNCIADO."""
    template = tmp_path / "plantilla.docx"
    source = tmp_path / "reporte.docx"
    _build_template(template, _LABELS)
    # En el Word del profesional la actividad 1 es la que en la plantilla es la
    # 3, y la 2 es la 1: mismos enunciados, numerales cruzados.
    _build_source(
        source,
        [
            (_LABELS[2], ["Concepto técnico del predio X"]),
            (_LABELS[0], ["Análisis de requisitos del proyecto Y"]),
        ],
    )

    use_case = _make_use_case(FakeProgress())
    plan = use_case.plan(_request(template, source, tmp_path, "salida.docx")).unwrap()

    assert plan.read_errors == ()
    by_ordinal = {p.ordinal: p for p in plan.planned if p.matched}
    assert set(by_ordinal) == {1, 3}  # numerales de la PLANTILLA, no del origen

    result = use_case.apply(
        _request(template, source, tmp_path, "salida.docx"), plan, {}
    ).unwrap()
    assert result.items_written == 2

    texts = _activity_texts(tmp_path / "salida.docx")
    assert "Concepto técnico del predio X" in texts[3]
    assert "Análisis de requisitos del proyecto Y" in texts[1]
    # Y no al revés (no se insertó por numeral).
    assert "Concepto técnico del predio X" not in texts[1]


def test_word_source_unmatched_label_goes_to_review(tmp_path: Path) -> None:
    template = tmp_path / "plantilla.docx"
    source = tmp_path / "reporte.docx"
    _build_template(template, _LABELS)
    _build_source(
        source,
        [
            (_LABELS[0], ["Contenido válido"]),
            ("Actividad de otro contrato totalmente distinta", ["Contenido huérfano"]),
        ],
    )

    use_case = _make_use_case(FakeProgress())
    plan = use_case.plan(_request(template, source, tmp_path, "salida.docx")).unwrap()

    unmatched = plan.unmatched
    assert len(unmatched) == 1
    assert "otro contrato" in unmatched[0].label
    # El numeral sintético NO coincide con ninguno de la plantilla.
    assert unmatched[0].ordinal not in {o.ordinal for o in plan.word_activities}


def test_word_source_no_activity_phrase_gets_default_text(tmp_path: Path) -> None:
    """Frases de "no se requirió" en el Word del profesional se filtran."""
    template = tmp_path / "plantilla.docx"
    source = tmp_path / "reporte.docx"
    _build_template(template, _LABELS[:2])
    _build_source(
        source,
        [
            (_LABELS[0], ["No se requirió esta actividad en el periodo"]),
            (_LABELS[1], ["Capacitación dictada en marzo"]),
        ],
    )

    use_case = _make_use_case(FakeProgress())
    result = use_case.execute(
        _request(template, source, tmp_path, "salida.docx")
    ).unwrap()

    assert result.items_written == 1
    texts = _activity_texts(tmp_path / "salida.docx")
    assert DEFAULT_EMPTY_ACTIVITY_TEXT in texts[1]
    assert "Capacitación dictada en marzo" in texts[2]


def test_template_added_as_source_is_skipped(tmp_path: Path) -> None:
    """Si la propia plantilla entra como archivo de profesional, se omite."""
    template = tmp_path / "plantilla.docx"
    _build_template(template, _LABELS)

    use_case = _make_use_case(FakeProgress())
    plan = use_case.plan(_request(template, template, tmp_path, "salida.docx")).unwrap()

    assert plan.professionals == ()
    assert any("misma plantilla" in w for w in plan.general_warnings)


def test_mixed_excel_and_word_sources(tmp_path: Path) -> None:
    """Un lote mixto (Excel + Word) se procesa con un solo flujo."""
    from tests.integration.test_cross_ods_validation import _build_excel

    template = tmp_path / "plantilla.docx"
    word_source = tmp_path / "reporte_word.docx"
    excel_source = tmp_path / "reporte_excel.xlsx"
    _build_template(template, _LABELS)
    _build_source(word_source, [(_LABELS[0], ["Aporte desde Word"])])
    _build_excel(excel_source, _LABELS)

    use_case = _make_use_case(FakeProgress())
    request = ProcessRequest(
        word_template=template,
        excel_files=(word_source, excel_source),
        output_dir=tmp_path,
        month="MARZO",
        output_name="mixto.docx",
    )
    result = use_case.execute(request).unwrap()

    assert result.professionals_processed == 2
    texts = _activity_texts(tmp_path / "mixto.docx")
    assert "Aporte desde Word" in texts[1]
    assert "Se realizó la actividad" in texts[1]  # contenido del Excel
