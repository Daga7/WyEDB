"""Pruebas de integración del lector de Word contra el documento real."""

from __future__ import annotations

from pathlib import Path

import docx
import pytest

from ods_reporter.infrastructure.word.docx_reader import DocxReader


@pytest.fixture
def reader() -> DocxReader:
    return DocxReader()


def test_reads_all_30_activities(reader: DocxReader, word_fixture: Path) -> None:
    document = docx.Document(str(word_fixture))
    activities = reader.read_activities(document)
    ordinals = [a.ordinal for a in activities]
    assert ordinals == list(range(1, 31))


def test_multirow_activity_has_multiple_entregables(
    reader: DocxReader, word_fixture: Path
) -> None:
    document = docx.Document(str(word_fixture))
    activities = reader.read_activities(document)
    # La actividad VI abarca 2 sub-filas (estructura limpia de plantilla) -> 2 entregables.
    act6 = next(a for a in activities if a.ordinal == 6)
    assert len(act6.entregables) == 2
    textos = [e.normalized_text for e in act6.entregables]
    assert any("informe del desarrollo" in t for t in textos)
    assert any("informes tecnicos" in t for t in textos)


def test_empty_activity_has_blank_template_slot(
    reader: DocxReader, word_fixture: Path
) -> None:
    # Una actividad no diligenciada conserva el slot vacío de la plantilla.
    document = docx.Document(str(word_fixture))
    activities = reader.read_activities(document)
    act2 = next(a for a in activities if a.ordinal == 2)
    entregable = act2.entregables[0]
    assert entregable.slot_paragraph is not None
    assert entregable.slot_paragraph.text.strip() == ""


def test_each_entregable_has_insertion_slot(reader: DocxReader, word_fixture: Path) -> None:
    document = docx.Document(str(word_fixture))
    activities = reader.read_activities(document)
    for activity in activities:
        for entregable in activity.entregables:
            assert entregable.slot_paragraph is not None
            assert entregable.cell is not None


def test_single_row_activity_has_one_entregable(
    reader: DocxReader, word_fixture: Path
) -> None:
    document = docx.Document(str(word_fixture))
    activities = reader.read_activities(document)
    act1 = next(a for a in activities if a.ordinal == 1)
    assert len(act1.entregables) == 1


def test_blank_template_reads_entregable_without_realizadas_header(
    reader: DocxReader, word_blank_fixture: Path
) -> None:
    # En la plantilla en blanco, la 1ª sub-fila de la actividad 4 NO trae el
    # encabezado "...realizadas:"; aun así debe leerse su entregable y su slot.
    document = docx.Document(str(word_blank_fixture))
    activities = reader.read_activities(document)
    act4 = next(a for a in activities if a.ordinal == 4)
    assert len(act4.entregables) == 2
    textos = [e.normalized_text for e in act4.entregables]
    assert any("cargar en repositorio" in t for t in textos)
    assert any("plan de accion actualizado" in t for t in textos)
