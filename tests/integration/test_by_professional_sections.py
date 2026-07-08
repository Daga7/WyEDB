"""Pruebas de plantillas divididas por profesional (tipo ODS 17).

El Word repite la lista de numerales por cada profesional, bajo una fila-título
"Actividades reportadas por el profesional X". El contenido de cada profesional
debe caer SOLO en su sección, aunque los numerales se repitan entre secciones.
"""

from __future__ import annotations

from pathlib import Path

import docx
import pytest

from ods_reporter.domain.entities.activity import Activity
from ods_reporter.domain.entities.entregable import Entregable
from ods_reporter.domain.value_objects.content_item import ContentItem
from ods_reporter.infrastructure.word.docx_processor import DocxProcessor
from ods_reporter.infrastructure.word.docx_reader import DocxReader

FIXTURES = Path(__file__).parent.parent / "fixtures"
BY_PROF = FIXTURES / "ODS17_by_professional.docx"


@pytest.fixture
def by_prof_template() -> Path:
    if not BY_PROF.exists():
        pytest.skip(f"Falta el fixture: {BY_PROF}")
    return BY_PROF


def _activity(ordinal: int, label: str, content: str) -> Activity:
    entregable = Entregable(
        entregable_text=label, content_items=(ContentItem(content),)
    )
    return Activity(ordinal=ordinal, label=label, entregables=(entregable,))


def test_reader_assigns_each_section_its_name(by_prof_template: Path) -> None:
    structure = DocxReader().read_structure(docx.Document(str(by_prof_template)))
    groups = {a.group_index for a in structure.activities}
    assert groups == {1, 2}
    g1 = [a for a in structure.activities if a.group_index == 1]
    g2 = [a for a in structure.activities if a.group_index == 2]
    assert all(a.professional_name == "Ana Perez" for a in g1)
    assert all(a.professional_name == "Beto Gomez" for a in g2)
    # Los numerales se repiten entre secciones.
    assert [a.ordinal for a in g1] == [1, 2]
    assert [a.ordinal for a in g2] == [1, 2]


def test_content_lands_in_matching_section(
    by_prof_template: Path, tmp_path: Path
) -> None:
    out = tmp_path / "salida.docx"
    processor = DocxProcessor()
    import shutil

    shutil.copy(by_prof_template, out)
    processor.open(out)

    # Se inserta primero el segundo profesional para comprobar que NO se coloca en
    # la primera sección por orden, sino en la que coincide por nombre.
    processor.insert_activity_content(
        _activity(1, "Primera actividad", "TAREA-DE-BETO"),
        professional_name="Beto Gomez Lopez",
    )
    processor.insert_activity_content(
        _activity(1, "Primera actividad", "TAREA-DE-ANA"),
        professional_name="Ana Maria Perez",
    )
    processor.save(out)

    structure = DocxReader().read_structure(docx.Document(str(out)))
    by_group = {}
    for a in structure.activities:
        if a.ordinal != 1:
            continue
        text = " ".join(e.slot_paragraph.text for e in a.entregables)
        by_group[a.professional_name] = text

    assert "TAREA-DE-ANA" in by_group["Ana Perez"]
    assert "TAREA-DE-BETO" in by_group["Beto Gomez"]
    # Y no se cruzan.
    assert "TAREA-DE-BETO" not in by_group["Ana Perez"]
    assert "TAREA-DE-ANA" not in by_group["Beto Gomez"]


def test_unknown_professional_uses_first_free_section_with_warning(
    by_prof_template: Path, tmp_path: Path
) -> None:
    out = tmp_path / "salida.docx"
    processor = DocxProcessor()
    import shutil

    shutil.copy(by_prof_template, out)
    processor.open(out)

    result = processor.insert_activity_content(
        _activity(1, "Primera actividad", "TAREA-DESCONOCIDA"),
        professional_name="Zoraida Quintero",
    )
    processor.save(out)

    assert any("no coincide" in w for w in result.warnings)
    structure = DocxReader().read_structure(docx.Document(str(out)))
    first = next(a for a in structure.activities if a.group_index == 1 and a.ordinal == 1)
    text = " ".join(e.slot_paragraph.text for e in first.entregables)
    assert "TAREA-DESCONOCIDA" in text
