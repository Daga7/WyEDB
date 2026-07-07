"""Pruebas del plan (sin escritura) y la reasignación manual en DocxProcessor.

Usa un documento Word sintético construido con python-docx, de modo que estas
pruebas corren también donde no están los fixtures reales (p. ej. en CI).
"""

from __future__ import annotations

from pathlib import Path

import docx
import pytest

from ods_reporter.domain.entities.activity import Activity
from ods_reporter.domain.entities.entregable import Entregable
from ods_reporter.domain.value_objects.content_item import ContentItem
from ods_reporter.infrastructure.word.docx_processor import DocxProcessor
from ods_reporter.infrastructure.word.docx_reader import DocxReader

_ACTIVITIES = (
    ("I", "Asesoría ambiental en campo"),
    ("II", "Capacitación en manejo de residuos"),
)


def _build_word(path: Path, *, with_observaciones: bool = False) -> None:
    """Construye una plantilla mínima con la estructura que espera el lector."""
    document = docx.Document()
    extra_rows = 1 if with_observaciones else 0
    table = document.add_table(rows=1 + len(_ACTIVITIES) + extra_rows, cols=2)
    table.rows[0].cells[0].text = "No"
    table.rows[0].cells[1].text = "Actividades"

    for index, (numeral, title) in enumerate(_ACTIVITIES, start=1):
        cells = table.rows[index].cells
        cells[0].text = numeral
        cell = cells[1]
        cell.paragraphs[0].text = f"Actividad: {title}"
        cell.add_paragraph(f"Descripción del entregable: Informe de {title.lower()}")
        cell.add_paragraph("Descripción de las actividades realizadas:")
        cell.add_paragraph("", style="List Bullet")  # slot de inserción

    if with_observaciones:
        # Fila final: la celda de texto es la PRIMERA de la fila (como en la
        # plantilla real, donde va combinada) con título + viñeta de inserción.
        cell = table.rows[-1].cells[0]
        cell.paragraphs[0].text = (
            "Observaciones generales y/o actividades adicionales encomendadas."
        )
        cell.add_paragraph("", style="List Bullet")

    document.save(str(path))


def _activity(ordinal: int, label: str, entregable: str, items: tuple[str, ...]) -> Activity:
    return Activity(
        ordinal=ordinal,
        label=label,
        entregables=(
            Entregable(
                entregable_text=entregable,
                content_items=tuple(ContentItem(text) for text in items),
            ),
        ),
    )


@pytest.fixture
def word_path(tmp_path: Path) -> Path:
    path = tmp_path / "plantilla.docx"
    _build_word(path)
    return path


def _full_text(path: Path) -> str:
    document = docx.Document(str(path))
    return "\n".join(
        p.text for t in document.tables for row in t.rows for c in row.cells for p in c.paragraphs
    )


def test_overview_lists_activities(word_path: Path) -> None:
    processor = DocxProcessor()
    processor.open(word_path)
    overview = processor.get_activities_overview()
    assert [(o.ordinal, o.entregable_count) for o in overview] == [(1, 1), (2, 1)]
    assert "Asesoría ambiental" in overview[0].label


def test_plan_matches_without_writing(word_path: Path, tmp_path: Path) -> None:
    processor = DocxProcessor()
    processor.open(word_path)
    activity = _activity(1, "Asesoría ambiental en campo", "Informe", ("Se asesoró la obra A",))

    outcome = processor.plan_activity_content(activity)

    assert outcome.matched is True
    assert outcome.items_written == 1
    assert outcome.warnings == ()  # el enunciado coincide: sin advertencias
    # El plan NO escribe: al guardar, el contenido no está en el documento.
    saved = tmp_path / "tras_plan.docx"
    processor.save(saved)
    assert "Se asesoró la obra A" not in _full_text(saved)


def test_same_ordinal_but_different_label_warns(word_path: Path) -> None:
    processor = DocxProcessor()
    processor.open(word_path)
    activity = _activity(
        1, "Totalmente otra actividad de otro proyecto", "Informe", ("Algo",)
    )

    outcome = processor.plan_activity_content(activity)

    assert outcome.matched is True  # el numeral existe...
    assert any("no coincide" in w for w in outcome.warnings)  # ...pero se advierte


def test_plan_reports_unmatched_activity(word_path: Path) -> None:
    processor = DocxProcessor()
    processor.open(word_path)
    activity = _activity(9, "Actividad inexistente", "Informe", ("Contenido huérfano",))

    outcome = processor.plan_activity_content(activity)

    assert outcome.matched is False
    assert any("no existe en el Word" in w for w in outcome.warnings)


def test_observaciones_section_detected_and_filled(tmp_path: Path) -> None:
    path = tmp_path / "con_observaciones.docx"
    _build_word(path, with_observaciones=True)

    processor = DocxProcessor()
    processor.open(path)

    # La fila de observaciones no cuenta como actividad.
    assert len(processor.get_activities_overview()) == len(_ACTIVITIES)
    assert processor.has_other_activities_section() is True

    items = (ContentItem("Capacitación extra (16/03/2026)"), ContentItem("Apoyo en comité"))
    assert processor.insert_other_activities(items) == 2

    saved = tmp_path / "con_observaciones_lleno.docx"
    processor.save(saved)
    assert "Capacitación extra (16/03/2026)" in _full_text(saved)
    assert "Apoyo en comité" in _full_text(saved)


def test_word_without_observaciones_reports_no_section(word_path: Path) -> None:
    processor = DocxProcessor()
    processor.open(word_path)
    assert processor.has_other_activities_section() is False
    assert processor.insert_other_activities((ContentItem("Algo"),)) == 0


def test_manual_reassignment_inserts_in_chosen_activity(
    word_path: Path, tmp_path: Path
) -> None:
    processor = DocxProcessor()
    processor.open(word_path)
    activity = _activity(9, "Actividad inexistente", "Informe", ("Contenido reasignado",))

    outcome = processor.insert_activity_content(activity, target_ordinal=2)

    assert outcome.matched is True
    assert outcome.items_written == 1
    assert any("manualmente" in w for w in outcome.warnings)

    saved = tmp_path / "reasignado.docx"
    processor.save(saved)

    # El contenido quedó en la actividad II (y solo allí).
    reread = docx.Document(str(saved))
    activities = DocxReader().read_activities(reread)
    by_ordinal = {a.ordinal: a for a in activities}
    text_act2 = "\n".join(p.text for p in by_ordinal[2].entregables[0].cell.paragraphs)
    text_act1 = "\n".join(p.text for p in by_ordinal[1].entregables[0].cell.paragraphs)
    assert "Contenido reasignado" in text_act2
    assert "Contenido reasignado" not in text_act1
