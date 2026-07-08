"""Pruebas de la localización FLEXIBLE de la tabla de actividades.

Documentos creados o guardados con distintas versiones de Word varían en
encabezados ("No", "N°", "No.", "Item"), filas de título encima del encabezado,
columnas adicionales, tablas previas o tablas anidadas. El lector debe
encontrar la tabla por su contenido en todos esos casos.
"""

from __future__ import annotations

from pathlib import Path

import docx
import pytest

from ods_reporter.domain.exceptions import WordProcessError
from ods_reporter.infrastructure.word.docx_processor import DocxProcessor
from ods_reporter.infrastructure.word.docx_reader import DocxReader

_ACTIVITIES = (
    ("I", "Asesoría ambiental en campo"),
    ("II", "Capacitación en manejo de residuos"),
)


def _fill_activity_rows(table, *, first_row: int, lead_cols: int) -> None:
    for index, (numeral, title) in enumerate(_ACTIVITIES):
        cells = table.rows[first_row + index].cells
        cells[lead_cols].text = numeral
        cell = cells[lead_cols + 1]
        cell.paragraphs[0].text = f"Actividad: {title}"
        cell.add_paragraph(f"Descripción del entregable: Informe de {title.lower()}")
        cell.add_paragraph("Descripción de las actividades realizadas:")
        cell.add_paragraph("", style="List Bullet")


def _build_word(
    path: Path,
    *,
    no_header: str = "No",
    activities_header: str = "Actividades",
    title_rows: int = 0,
    lead_cols: int = 0,
    decoy_table: bool = False,
    nested: bool = False,
) -> None:
    """Construye una plantilla con variaciones estructurales controladas."""
    document = docx.Document()
    if decoy_table:
        decoy = document.add_table(rows=2, cols=2)
        decoy.rows[0].cells[0].text = "Elaborado por"
        decoy.rows[0].cells[1].text = "Fecha"

    total_rows = title_rows + 1 + len(_ACTIVITIES)
    cols = 2 + lead_cols
    if nested:
        host = document.add_table(rows=1, cols=1)
        table = host.rows[0].cells[0].add_table(rows=total_rows, cols=cols)
    else:
        table = document.add_table(rows=total_rows, cols=cols)

    if title_rows:
        table.rows[0].cells[0].text = "INFORME MENSUAL DE ACTIVIDADES — ODS"
    header = table.rows[title_rows].cells
    header[lead_cols].text = no_header
    header[lead_cols + 1].text = activities_header
    if lead_cols:
        header[0].text = "Ítem previo"

    _fill_activity_rows(table, first_row=title_rows + 1, lead_cols=lead_cols)
    document.save(str(path))


def _assert_structure_ok(path: Path) -> None:
    document = docx.Document(str(path))
    structure = DocxReader().read_structure(document)
    assert [a.ordinal for a in structure.activities] == [1, 2]
    assert "Asesoría ambiental" in structure.activities[0].label
    assert all(len(a.entregables) == 1 for a in structure.activities)


@pytest.mark.parametrize(
    ("no_header", "activities_header"),
    [
        ("No", "Actividades"),  # plantilla clásica (compatibilidad intacta)
        ("N°", "ACTIVIDAD"),
        ("No.", "Actividades a desarrollar"),
        ("Nro.", "ACTIVIDADES"),
        ("ITEM", "Actividades"),
        ("Nº", "Actividades del periodo"),
    ],
)
def test_header_variants_are_recognized(
    tmp_path: Path, no_header: str, activities_header: str
) -> None:
    path = tmp_path / "variante.docx"
    _build_word(path, no_header=no_header, activities_header=activities_header)
    _assert_structure_ok(path)


def test_title_rows_above_header(tmp_path: Path) -> None:
    path = tmp_path / "con_titulo.docx"
    _build_word(path, title_rows=2)
    _assert_structure_ok(path)


def test_extra_leading_column(tmp_path: Path) -> None:
    path = tmp_path / "columna_extra.docx"
    _build_word(path, lead_cols=1)
    _assert_structure_ok(path)


def test_decoy_table_before_activities(tmp_path: Path) -> None:
    path = tmp_path / "con_tabla_previa.docx"
    _build_word(path, decoy_table=True)
    _assert_structure_ok(path)


def test_nested_activities_table(tmp_path: Path) -> None:
    path = tmp_path / "anidada.docx"
    _build_word(path, nested=True)
    _assert_structure_ok(path)


def test_all_variations_combined(tmp_path: Path) -> None:
    path = tmp_path / "combinada.docx"
    _build_word(
        path,
        no_header="N°",
        activities_header="ACTIVIDADES A DESARROLLAR",
        title_rows=1,
        lead_cols=1,
        decoy_table=True,
    )
    _assert_structure_ok(path)


def test_processor_opens_variant_template(tmp_path: Path) -> None:
    """El procesador completo (overview) también funciona con las variantes."""
    path = tmp_path / "variante_procesador.docx"
    _build_word(path, no_header="N°", activities_header="ACTIVIDAD", title_rows=1)
    processor = DocxProcessor()
    processor.open(path)
    overview = processor.get_activities_overview()
    assert [(o.ordinal, o.entregable_count) for o in overview] == [(1, 1), (2, 1)]


def test_numbered_decoy_table_without_slots_is_not_chosen(tmp_path: Path) -> None:
    """Un cronograma/índice numerado ANTES de la tabla real no debe ganarle.

    El cronograma también tiene encabezado tipo 'No'/'Actividades' y filas
    numeradas, pero sus celdas no tienen slots de inserción (viñetas de
    plantilla): la tabla real se reconoce por tenerlos.
    """
    path = tmp_path / "con_cronograma.docx"
    document = docx.Document()

    decoy = document.add_table(rows=3, cols=2)
    decoy.rows[0].cells[0].text = "No"
    decoy.rows[0].cells[1].text = "Actividades programadas"
    decoy.rows[1].cells[0].text = "1"
    decoy.rows[1].cells[1].text = "Semana 1: inducción del proyecto"
    decoy.rows[2].cells[0].text = "2"
    decoy.rows[2].cells[1].text = "Semana 2: entrega parcial"

    table = document.add_table(rows=1 + len(_ACTIVITIES), cols=2)
    table.rows[0].cells[0].text = "No"
    table.rows[0].cells[1].text = "Actividades"
    _fill_activity_rows(table, first_row=1, lead_cols=0)
    document.save(str(path))

    structure = DocxReader().read_structure(docx.Document(str(path)))

    # Se eligió la tabla REAL: enunciados correctos y con slots de inserción.
    assert "Asesoría ambiental" in structure.activities[0].label
    assert all(a.entregables for a in structure.activities)
    assert not any("Semana" in a.label for a in structure.activities)


def _build_ecopetrol_style(path: Path) -> None:
    """Plantilla tipo ECOPETROL (FO-GT-EP-093): rótulo "Alcance específico:",
    columnas de porcentaje y numerales romanos con punto ("VI.", "XXI.\xa0")."""
    document = docx.Document()
    labels = [
        ("I", "Coordinación y seguimiento a la ejecución técnica"),
        ("VI.", "Dimensionar necesidades de recursos y equipos"),
        ("XXI.\xa0", "Seguimiento a la facturación, pagos y causación"),
    ]
    table = document.add_table(rows=1 + len(labels), cols=4)
    header = table.rows[0].cells
    header[0].text = "N.º"
    header[1].text = "Descripción de Actividad"
    header[2].text = "% Plan"
    header[3].text = "% Real"
    for index, (numeral, title) in enumerate(labels, start=1):
        cells = table.rows[index].cells
        cells[0].text = numeral
        cell = cells[1]
        cell.paragraphs[0].text = "Alcance específico:"
        cell.add_paragraph(f"\xa0\xa0 {['i', 'vi', 'xxi'][index - 1]}.\xa0\xa0 {title}")
        cell.add_paragraph("Descripción del entregable: ")
        cell.add_paragraph(f"Entregable de {title.lower()}")
        cell.add_paragraph("Descripción de las actividades realizadas:")
        cell.add_paragraph("", style="List Bullet")
        cells[2].text = "82%"
        cells[3].text = "82%"
    document.save(str(path))


def test_ecopetrol_style_template(tmp_path: Path) -> None:
    """Numerales con punto y rótulo 'Alcance específico:' se leen bien.

    Regresión del caso real ODS 13: antes, "VI." perdía su numeral (la fila se
    pegaba a la actividad anterior) y el enunciado salía como el rótulo
    "Alcance específico:", con lo que la validación rechazaba archivos buenos.
    """
    path = tmp_path / "ecopetrol.docx"
    _build_ecopetrol_style(path)

    structure = DocxReader().read_structure(docx.Document(str(path)))

    assert [a.ordinal for a in structure.activities] == [1, 6, 21]
    labels = [a.label for a in structure.activities]
    assert labels[0].startswith("Coordinación y seguimiento")
    assert labels[1].startswith("Dimensionar necesidades")
    assert labels[2].startswith("Seguimiento a la facturación")
    assert not any("Alcance" in label for label in labels)
    assert all(len(a.entregables) == 1 for a in structure.activities)
    # El entregable también se leyó con su marcador clásico.
    assert structure.activities[0].entregables[0].entregable_text.startswith("Entregable de")


def test_real_ods13_template_and_excel_are_compatible(tmp_path: Path) -> None:
    """Con los archivos reales de la ODS 13, la validación acepta el Excel."""
    word = Path("tests/fixtures/ODS13_Word.docx")
    excel = Path("tests/fixtures/ODS13_Excel.xlsx")
    if not word.exists() or not excel.exists():
        pytest.skip("Faltan los fixtures reales de la ODS 13")

    from tests.integration.test_process_ods import FakeProgress, _make_use_case
    from ods_reporter.application.use_cases.process_ods import ProcessRequest

    use_case = _make_use_case(FakeProgress())
    request = ProcessRequest(
        word_template=word,
        excel_files=(excel,),
        output_dir=tmp_path,
        month="ENERO",
        output_name="ods13.docx",
    )
    plan = use_case.plan(request).unwrap()

    assert plan.read_errors == ()
    assert len(plan.professionals) == 1
    assert len(plan.word_activities) == 21
    assert plan.unmatched == ()


def test_document_without_activities_table_raises(tmp_path: Path) -> None:
    path = tmp_path / "sin_tabla.docx"
    document = docx.Document()
    table = document.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "Nombre"
    table.rows[0].cells[1].text = "Cargo"
    document.save(str(path))

    with pytest.raises(WordProcessError, match="No se encontró la tabla de actividades"):
        DocxReader().read_structure(docx.Document(str(path)))
