"""Pruebas del lector de Word con DISTINTAS plantillas (genérico).

Verifica que el mismo lector reconoce las actividades en plantillas con formatos
diferentes: ODS 11 (numeración romana, con marcadores) y ODS 266 (numeración
arábiga, sin marcadores).
"""

from __future__ import annotations

from pathlib import Path

import docx
import pytest

from ods_reporter.infrastructure.word.docx_reader import DocxReader

FIXTURES = Path(__file__).parent.parent / "fixtures"
ODS266 = FIXTURES / "ODS266_Word.docx"
ODS18_LIKE = FIXTURES / "ODS18_like_cell.docx"


@pytest.fixture
def reader() -> DocxReader:
    return DocxReader()


@pytest.fixture
def ods266() -> Path:
    if not ODS266.exists():
        pytest.skip(f"Falta el fixture: {ODS266}")
    return ODS266


def test_reads_arabic_numbered_template(reader: DocxReader, ods266: Path) -> None:
    activities = reader.read_activities(docx.Document(str(ods266)))
    # Numeración arábiga 1..46, una actividad por fila.
    ordinals = [a.ordinal for a in activities]
    assert ordinals == list(range(1, 47))


def test_markerless_cells_have_one_entregable_with_slot(
    reader: DocxReader, ods266: Path
) -> None:
    activities = reader.read_activities(docx.Document(str(ods266)))
    for activity in activities:
        assert len(activity.entregables) == 1
        assert activity.entregables[0].slot_paragraph is not None


def test_markerless_entregable_uses_activity_title(reader: DocxReader, ods266: Path) -> None:
    activities = reader.read_activities(docx.Document(str(ods266)))
    first = activities[0]
    # Sin marcador "Descripción del entregable:", el texto es el título.
    assert "informe" in first.entregables[0].normalized_text


def test_section_rows_are_skipped(reader: DocxReader, ods266: Path) -> None:
    # Las filas de sección (que abarcan las 4 columnas) no generan actividades.
    activities = reader.read_activities(docx.Document(str(ods266)))
    # 46 actividades exactas pese a las filas de sección intercaladas.
    assert len(activities) == 46


# --- Plantilla tipo ODS 18: rótulo "Actividades:" y "Descripción entregable:" ---


@pytest.fixture
def ods18_like() -> Path:
    if not ODS18_LIKE.exists():
        pytest.skip(f"Falta el fixture: {ODS18_LIKE}")
    return ODS18_LIKE


def test_actividades_plural_marker_reads_real_enunciado(
    reader: DocxReader, ods18_like: Path
) -> None:
    """El enunciado va bajo el rótulo plural 'Actividades:', no en el rótulo mismo.

    Antes, el lector devolvía la palabra literal 'Actividades:' como enunciado, lo
    que hacía que la comparación con el Excel diera 'enunciado distinto' aunque
    fueran idénticos (bug reportado en la ODS 18).
    """
    activities = reader.read_activities(docx.Document(str(ods18_like)))
    assert len(activities) == 1
    label = activities[0].label
    assert "Elaboración de diagnóstico ambiental" in label
    assert "Actividades:" not in label


def test_entregable_marker_without_del_is_boundary(
    reader: DocxReader, ods18_like: Path
) -> None:
    """'Descripción entregable:' (sin 'del') cierra el enunciado igual que la variante con 'del'."""
    activities = reader.read_activities(docx.Document(str(ods18_like)))
    label = activities[0].label
    # El texto del entregable NO debe haberse arrastrado dentro del enunciado.
    assert "Repositorio" not in label
    # Y el entregable se lee por separado.
    assert "Repositorio" in activities[0].entregables[0].entregable_text
