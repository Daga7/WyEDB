"""Pruebas de integración del escritor de Word: inserción preservando formato."""

from __future__ import annotations

from pathlib import Path

import docx
import pytest

from ods_reporter.domain.value_objects.content_item import ContentItem
from ods_reporter.infrastructure.word.docx_reader import DocxReader
from ods_reporter.infrastructure.word.docx_writer import DocxWriter


@pytest.fixture
def reader() -> DocxReader:
    return DocxReader()


@pytest.fixture
def writer() -> DocxWriter:
    return DocxWriter()


def _activity(activities, ordinal):  # type: ignore[no-untyped-def]
    return next(a for a in activities if a.ordinal == ordinal)


def test_fill_writes_items_and_preserves_style(
    reader: DocxReader, writer: DocxWriter, word_fixture: Path, tmp_path: Path
) -> None:
    document = docx.Document(str(word_fixture))
    activities = reader.read_activities(document)
    # Actividad 2: vacía -> slot de plantilla limpio.
    entregable = _activity(activities, 2).entregables[0]
    slot_style = entregable.slot_paragraph.style.name  # "List Paragraph"

    items = (ContentItem("Primer ítem"), ContentItem("Segundo ítem"), ContentItem("Tercer ítem"))
    written = writer.fill_entregable(entregable, items)
    assert written == 3

    # Guardar y recargar para confirmar persistencia.
    out = tmp_path / "salida.docx"
    document.save(str(out))
    reloaded = docx.Document(str(out))
    activities2 = reader.read_activities(reloaded)
    cell = _activity(activities2, 2).entregables[0].cell

    # Los 3 ítems están presentes.
    cell_text = cell.text
    for esperado in ("Primer ítem", "Segundo ítem", "Tercer ítem"):
        assert esperado in cell_text

    # Los párrafos insertados conservan el estilo del slot original (List Paragraph).
    inserted = [p for p in cell.paragraphs if p.text.strip() and p.style.name == slot_style]
    assert len(inserted) == 3


def test_set_default_text_fills_empty_slot(
    reader: DocxReader, writer: DocxWriter, word_fixture: Path, tmp_path: Path
) -> None:
    document = docx.Document(str(word_fixture))
    activities = reader.read_activities(document)
    entregable = _activity(activities, 2).entregables[0]

    writer.set_default_text(entregable, "Durante el periodo se estuvo atento.")

    out = tmp_path / "salida.docx"
    document.save(str(out))
    reloaded = docx.Document(str(out))
    cell = _activity(reader.read_activities(reloaded), 2).entregables[0].cell
    assert "Durante el periodo se estuvo atento." in cell.text


def test_does_not_disturb_other_activities(
    reader: DocxReader, writer: DocxWriter, word_fixture: Path, tmp_path: Path
) -> None:
    document = docx.Document(str(word_fixture))
    activities = reader.read_activities(document)
    # Contenido existente de la actividad 1 (ya diligenciada en abril).
    act1_text_before = _activity(activities, 1).entregables[0].cell.text

    # Se rellena la actividad 2.
    writer.fill_entregable(_activity(activities, 2).entregables[0], (ContentItem("X"),))

    out = tmp_path / "salida.docx"
    document.save(str(out))
    reloaded = docx.Document(str(out))
    act1_text_after = _activity(reader.read_activities(reloaded), 1).entregables[0].cell.text
    assert act1_text_after == act1_text_before


def test_empty_items_writes_nothing(
    reader: DocxReader, writer: DocxWriter, word_fixture: Path
) -> None:
    document = docx.Document(str(word_fixture))
    activities = reader.read_activities(document)
    entregable = _activity(activities, 2).entregables[0]
    assert writer.fill_entregable(entregable, ()) == 0


def test_inserted_run_inherits_paragraph_mark_size(
    reader: DocxReader, writer: DocxWriter, word_fixture: Path, tmp_path: Path
) -> None:
    """El run insertado toma el tamaño/idioma de la marca de párrafo del slot.

    Sin esto, un run "pelado" heredaba el tamaño del estilo (a veces menor), que
    era la causa de que la fuente encogiera al copiar al Word.
    """
    from docx.oxml.ns import qn

    document = docx.Document(str(word_fixture))
    activities = reader.read_activities(document)
    entregable = _activity(activities, 2).entregables[0]

    # Tamaño previsto por la plantilla (marca de párrafo del slot).
    p_pr = entregable.slot_paragraph._p.find(qn("w:pPr"))
    mark_rpr = p_pr.find(qn("w:rPr"))
    mark_sz = mark_rpr.find(qn("w:sz")).get(qn("w:val"))

    writer.fill_entregable(entregable, (ContentItem("Contenido con tildes: acción"),))

    out = tmp_path / "salida.docx"
    document.save(str(out))
    reloaded = docx.Document(str(out))
    cell = _activity(reader.read_activities(reloaded), 2).entregables[0].cell
    run = next(
        run
        for paragraph in cell.paragraphs
        for run in paragraph.runs
        if run.text.strip()
    )
    run_rpr = run._element.find(qn("w:rPr"))
    run_sz = run_rpr.find(qn("w:sz")) if run_rpr is not None else None
    assert run_sz is not None
    assert run_sz.get(qn("w:val")) == mark_sz
