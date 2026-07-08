"""Lector de reportes de profesionales en formato Word (modo Word → Word).

Lee el documento Word diligenciado por un profesional —con la misma tabla de
actividades que la plantilla— y lo convierte al ``RawReport`` del puerto de
lectura. Así todo el flujo posterior (normalización, reglas de negocio,
validación anti "ODS equivocada", revisión y escritura) es exactamente el
mismo que en el modo Excel → Word: solo cambia la pieza que lee.
"""

from __future__ import annotations

import logging
from pathlib import Path
from typing import TYPE_CHECKING

import docx

from ods_reporter.application.ports.excel_reader_port import (
    RawActivity,
    RawEntregable,
    RawOtherActivity,
    RawReport,
)
from ods_reporter.domain.entities.ods_metadata import ODSMetadata
from ods_reporter.domain.exceptions import InvalidInputError, WordProcessError
from ods_reporter.infrastructure.word.docx_reader import (
    DocxReader,
    entregable_content_texts,
    find_ods_number,
)
from ods_reporter.shared.constants import WORD_EXTENSIONS
from ods_reporter.shared.text_utils import normalize_text

if TYPE_CHECKING:
    from docx.document import Document

logger = logging.getLogger(__name__)

# Etiquetas (normalizadas) tras las que suele venir el nombre del profesional.
_NAME_LABELS = (
    "profesional responsable",
    "nombre del profesional",
    "elaborado por",
    "profesional",
)


class DocxSourceReader:
    """Lee el Word de un profesional y lo entrega como ``RawReport``.

    Implementa el mismo contrato que el lector de Excel (``ExcelReaderPort``);
    el parámetro ``month`` se acepta por contrato y se ignora, porque un Word
    no tiene hojas por mes: el documento ES el reporte del periodo.
    """

    def __init__(self, reader: DocxReader | None = None) -> None:
        self._reader = reader or DocxReader()

    def read_month(self, file_path: Path, month: str) -> RawReport:
        if not file_path.exists():
            raise InvalidInputError(f"No existe el archivo: {file_path}")
        if file_path.suffix.lower() not in WORD_EXTENSIONS:
            raise InvalidInputError(f"No es un documento Word: {file_path.name}")
        try:
            document = docx.Document(str(file_path))
        except Exception as exc:
            raise WordProcessError(
                f"No se pudo abrir el Word '{file_path.name}': {exc}"
            ) from exc

        structure = self._reader.read_structure(document)
        activities = tuple(
            RawActivity(
                ordinal=activity.ordinal,
                label=activity.label,
                entregables=tuple(
                    RawEntregable(
                        entregable_text=entregable.entregable_text,
                        raw_content="\n".join(entregable_content_texts(entregable)),
                    )
                    for entregable in activity.entregables
                ),
            )
            for activity in structure.activities
        )

        other_activities: tuple[RawOtherActivity, ...] = ()
        if structure.observaciones is not None:
            other_activities = tuple(
                RawOtherActivity(text=text)
                for text in entregable_content_texts(structure.observaciones)
            )

        name = self._find_professional_name(document) or file_path.stem
        metadata = ODSMetadata(
            ods_number=find_ods_number(document),
            responsible_professional=name,
            source_file=file_path.name,
        )
        logger.info(
            "Word de profesional '%s' leído (%s): %d actividad(es).",
            file_path.name,
            name,
            len(activities),
        )
        return RawReport(
            metadata=metadata, activities=activities, other_activities=other_activities
        )

    # --- Nombre del profesional (mejor esfuerzo) ---

    @staticmethod
    def _find_professional_name(document: Document) -> str:
        """Busca "Profesional: X" / "Elaborado por: X" en la cabecera del doc.

        Si no aparece, el llamador usa el nombre del archivo (misma regla de
        respaldo que el lector de Excel).
        """
        for paragraph in document.paragraphs[:40]:
            name = _name_after_label(paragraph.text)
            if name:
                return name
        for table in document.tables[:2]:
            for row in table.rows[:15]:
                cells = row.cells
                for index, cell in enumerate(cells):
                    name = _name_after_label(cell.text)
                    if name:
                        return name
                    # Etiqueta en una celda y el valor en la siguiente.
                    if normalize_text(cell.text) in _NAME_LABELS and index + 1 < len(cells):
                        value = cells[index + 1].text.strip()
                        if value:
                            return value
        return ""


def _name_after_label(text: str) -> str:
    """Extrae el valor de "Etiqueta: valor" si la etiqueta es de profesional."""
    normalized = normalize_text(text)
    for label in _NAME_LABELS:
        if normalized.startswith(label) and ":" in text:
            value = text.split(":", 1)[1].strip()
            if value:
                return value
    return ""
