"""Lector de la estructura de actividades de un documento Word de ODS.

Localiza la tabla de actividades y extrae, por cada actividad, sus entregables
(sub-filas) junto con la **ubicación exacta** donde se debe insertar el contenido
(la celda y el párrafo "Descripción de las actividades realizadas:").

Devuelve descriptores que conservan referencias vivas a los objetos de
python-docx, de modo que el escritor (``docx_writer``) pueda insertar después.
"""

from __future__ import annotations

import re
from dataclasses import dataclass, field
from typing import TYPE_CHECKING

from docx.oxml.ns import qn
from docx.table import Table, _Cell

from ods_reporter.domain.exceptions import WordProcessError
from ods_reporter.infrastructure.matching.roman_numerals import roman_to_int
from ods_reporter.shared.text_utils import (
    extract_ods_number,
    is_blank_or_placeholder,
    normalize_text,
)

if TYPE_CHECKING:
    from docx.document import Document
    from docx.text.paragraph import Paragraph

# Encabezado de la columna del numeral: variantes reales según la versión de
# Word / la persona que armó la plantilla ("No", "No.", "N°", "Nro", "Item"…).
# Se compara sobre el texto normalizado y sin signos (solo letras y dígitos).
_NO_HEADER_TOKENS = frozenset({"no", "n", "nro", "num", "numero", "item", "id"})
_NON_ALNUM_RE = re.compile(r"[^a-z0-9]+")

# Máximo de filas iniciales de una tabla donde se busca la fila de encabezado
# (tolera filas de título/logotipo encima del encabezado real).
_HEADER_SCAN_ROWS = 10

# Longitud máxima de un encabezado de columna de actividades: un texto largo que
# contiene "actividad" es contenido, no un encabezado.
_MAX_ACTIVITIES_HEADER_LEN = 60

# Marcadores de sección dentro de la celda de una actividad (normalizados).
_MARKER_ACTIVITY = "actividad:"
_MARKER_ENTREGABLE = "descripcion del entregable:"
_MARKER_REALIZADAS = "descripcion de las actividades realizadas:"

# Inicio (normalizado) de la fila de observaciones/actividades adicionales, que
# aparece tras la última actividad ("Observaciones generales y/o actividades
# adicionales encomendadas por...").
_MARKER_OBSERVACIONES = "observaciones"


@dataclass(slots=True)
class WordEntregable:
    """Un entregable del Word con la ubicación donde insertar su contenido.

    Attributes
    ----------
    entregable_text:
        Texto del entregable (para alinear con el Excel).
    cell:
        Celda de la tabla que contiene este entregable.
    slot_paragraph:
        Párrafo-plantilla (la viñeta ``List Paragraph`` del entregable) que se usa
        para clonar el formato e insertar el contenido.
    """

    entregable_text: str
    cell: _Cell
    slot_paragraph: Paragraph

    @property
    def normalized_text(self) -> str:
        return normalize_text(self.entregable_text)


@dataclass(slots=True)
class WordActivity:
    """Una actividad del Word: su numeral y sus entregables."""

    ordinal: int
    label: str
    entregables: list[WordEntregable] = field(default_factory=list)


@dataclass(slots=True)
class WordStructure:
    """Estructura completa del documento: actividades + sección de observaciones."""

    activities: list[WordActivity] = field(default_factory=list)
    observaciones: WordEntregable | None = None


@dataclass(slots=True)
class _TableMatch:
    """Una tabla candidata: dónde está su encabezado y en qué columnas."""

    table: Table
    header_row: int
    no_col: int
    activities_col: int


class DocxReader:
    """Extrae la estructura de actividades de un documento Word."""

    def read_structure(self, document: Document) -> WordStructure:
        """Devuelve las actividades y la sección de observaciones (si existe).

        La tabla de actividades se busca por CONTENIDO en todas las tablas del
        documento (incluidas las anidadas): una fila de encabezado con una
        columna de numeral ("No", "N°", "Item"…) y una de actividades, en
        cualquier posición. Si varias tablas parecen candidatas, gana la
        primera que realmente contenga actividades.

        Raises
        ------
        WordProcessError
            Si ninguna tabla del documento tiene esa forma.
        """
        candidates = self._find_table_candidates(document)
        if not candidates:
            raise WordProcessError(
                "No se encontró la tabla de actividades: ninguna tabla del "
                "documento tiene un encabezado tipo 'No'/'N°'/'Item' junto a "
                "'Actividades'. Verifique que el archivo sea un informe ODS."
            )
        first: WordStructure | None = None
        for candidate in candidates:
            structure = self._parse_table(candidate)
            if structure.activities:
                return structure
            if first is None:
                first = structure
        assert first is not None  # hay al menos una candidata
        return first

    def _parse_table(self, match: _TableMatch) -> WordStructure:
        table = match.table
        structure = WordStructure()
        current: WordActivity | None = None

        # Se itera a nivel XML (cada <w:tr> tiene sus propias <w:tc>): con celdas
        # combinadas verticalmente, ``row.cells`` de python-docx desplaza columnas.
        for tr in table._tbl.tr_lst[match.header_row + 1 :]:
            tcs = tr.tc_lst
            if not tcs:
                continue
            if len(tcs) < 2 or len(tcs) <= match.no_col:
                # Fila con las columnas combinadas en una sola celda: puede ser
                # la de observaciones/actividades adicionales.
                merged = _Cell(tcs[0], table)
                if self._is_observaciones_cell(merged):
                    structure.observaciones = self._read_observaciones(merged)
                    current = None
                continue
            col_no = _Cell(tcs[match.no_col], table)
            col_activity = _Cell(tcs[min(match.activities_col, len(tcs) - 1)], table)

            ordinal = self._read_ordinal(col_no)

            # Fila de observaciones/actividades adicionales: cierra la lista de
            # actividades (su celda de texto es la PRIMERA de la fila).
            if ordinal is None and self._is_observaciones_cell(col_no):
                structure.observaciones = self._read_observaciones(col_no)
                current = None
                continue

            entregable = self._read_entregable(col_activity)

            if ordinal is not None and (current is None or ordinal != current.ordinal):
                current = WordActivity(ordinal=ordinal, label=self._read_label(col_activity))
                structure.activities.append(current)

            if current is not None and entregable is not None:
                current.entregables.append(entregable)

        return structure

    def read_activities(self, document: Document) -> list[WordActivity]:
        """Devuelve solo las actividades (compatibilidad con el uso existente)."""
        return self.read_structure(document).activities

    # --- Sección de observaciones ---

    @staticmethod
    def _is_observaciones_cell(cell: _Cell) -> bool:
        return normalize_text(cell.text).startswith(_MARKER_OBSERVACIONES)

    def _read_observaciones(self, cell: _Cell) -> WordEntregable | None:
        """Extrae la sección de observaciones: título + viñeta de inserción."""
        slot = self._find_slot(cell.paragraphs, None)
        if slot is None:
            return None
        return WordEntregable(
            entregable_text=_first_title_text(cell.paragraphs), cell=cell, slot_paragraph=slot
        )

    # --- Localización de la tabla ---

    def _find_table_candidates(self, document: Document) -> list[_TableMatch]:
        """Tablas del documento (anidadas incluidas) con encabezado de actividades."""
        candidates: list[_TableMatch] = []
        for table in self._iter_tables(document):
            found = self._find_header(table)
            if found is not None:
                candidates.append(found)
        return candidates

    @staticmethod
    def _iter_tables(document: Document) -> list[Table]:
        """Todas las tablas del cuerpo en orden de documento, anidadas incluidas.

        ``document.tables`` solo devuelve las de primer nivel; iterar los
        elementos ``<w:tbl>`` del XML alcanza también las tablas dentro de
        celdas (documentos reorganizados por otras versiones de Word).
        """
        return [Table(tbl, document) for tbl in document.element.body.iter(qn("w:tbl"))]

    def _find_header(self, table: Table) -> _TableMatch | None:
        """Busca la fila de encabezado en las primeras filas de la tabla."""
        for row_idx, tr in enumerate(table._tbl.tr_lst[:_HEADER_SCAN_ROWS]):
            tcs = tr.tc_lst
            if len(tcs) < 2:
                continue
            texts = [normalize_text(_Cell(tc, table).text) for tc in tcs]
            no_col = next(
                (i for i, text in enumerate(texts) if self._is_no_header(text)), None
            )
            if no_col is None:
                continue
            activities_col = next(
                (
                    i
                    for i, text in enumerate(texts)
                    if i > no_col and self._is_activities_header(text)
                ),
                None,
            )
            if activities_col is None:
                continue
            return _TableMatch(
                table=table,
                header_row=row_idx,
                no_col=no_col,
                activities_col=activities_col,
            )
        return None

    @staticmethod
    def _is_no_header(normalized: str) -> bool:
        return _NON_ALNUM_RE.sub("", normalized) in _NO_HEADER_TOKENS

    @staticmethod
    def _is_activities_header(normalized: str) -> bool:
        return "actividad" in normalized and len(normalized) <= _MAX_ACTIVITIES_HEADER_LEN

    # --- Lectura de celdas ---

    @staticmethod
    def _read_ordinal(cell: _Cell) -> int | None:
        """Lee el numeral del 'No': admite romanos (I, II) y arábigos (1, 2)."""
        text = cell.text.strip()
        roman = roman_to_int(text)
        if roman is not None:
            return roman
        try:
            number = int(float(text))
        except (ValueError, TypeError):
            return None
        return number if number > 0 else None

    @staticmethod
    def _read_label(cell: _Cell) -> str:
        paragraphs = cell.paragraphs
        idx = _find_marker(paragraphs, _MARKER_ACTIVITY)
        if idx is not None:
            return _text_until_next_marker(paragraphs, idx)
        # Plantillas sin el marcador "Actividad:": el título es el primer texto.
        return _first_title_text(paragraphs)

    def _read_entregable(self, cell: _Cell) -> WordEntregable | None:
        """Extrae un entregable de la celda, soportando dos formatos de plantilla:

        - Con marcadores ("Descripción del entregable:" / "...realizadas:").
        - Sin marcadores: el título de la actividad y, debajo, las viñetas.

        En ambos casos el slot de inserción es la primera viñeta (``List Paragraph``).
        """
        paragraphs = cell.paragraphs
        entregable_idx = _find_marker(paragraphs, _MARKER_ENTREGABLE)
        realizadas_idx = _find_marker(paragraphs, _MARKER_REALIZADAS)

        slot = self._find_slot(paragraphs, realizadas_idx)
        if slot is None:
            return None  # sin viñeta no hay dónde insertar: no es celda de contenido

        if entregable_idx is not None:
            entregable_text = _text_until_next_marker(paragraphs, entregable_idx)
        else:
            # Sin marcador de entregable: se usa el título de la actividad como clave.
            entregable_text = _first_title_text(paragraphs)

        return WordEntregable(entregable_text=entregable_text, cell=cell, slot_paragraph=slot)

    @staticmethod
    def _find_slot(paragraphs: list[Paragraph], realizadas_idx: int | None) -> Paragraph | None:
        """Localiza el slot de inserción: la viñeta de plantilla (``List Paragraph``).

        Robusto ante plantillas donde, en algunas sub-filas, falta el encabezado
        "Descripción de las actividades realizadas:": en ese caso el slot es
        igualmente la primera viñeta que sigue al entregable.
        """
        start = realizadas_idx + 1 if realizadas_idx is not None else 0
        for paragraph in paragraphs[start:]:
            if _is_list_paragraph(paragraph):
                return paragraph
        # Respaldos: cualquier viñeta de la celda, o el párrafo tras el encabezado.
        for paragraph in paragraphs:
            if _is_list_paragraph(paragraph):
                return paragraph
        if realizadas_idx is not None and realizadas_idx + 1 < len(paragraphs):
            return paragraphs[realizadas_idx + 1]
        return None


# --- Utilidades sobre documentos ya leídos ---


def entregable_content_texts(entregable: WordEntregable) -> list[str]:
    """Textos de contenido ya escritos en un entregable (sus viñetas no vacías).

    Es la operación inversa del escritor: recupera lo que un profesional
    diligenció en su documento Word (modo Word → Word). Los slots vacíos o con
    solo un marcador de posición ("-", "•") se ignoran.
    """
    texts: list[str] = []
    for paragraph in entregable.cell.paragraphs:
        if not _is_list_paragraph(paragraph):
            continue
        text = paragraph.text.strip()
        if not text or is_blank_or_placeholder(text):
            continue
        texts.append(text)
    return texts


def find_ods_number(document: Document) -> str:
    """Busca el número de ODS en el texto del documento (mejor esfuerzo).

    Recorre los párrafos iniciales y las celdas de las primeras tablas (donde
    está la cabecera del informe, p. ej. "3040727 ECP ODS No. 11").
    """
    for paragraph in document.paragraphs[:40]:
        number = extract_ods_number(paragraph.text)
        if number:
            return number
    for table in document.tables[:2]:
        for row in table.rows[:15]:
            for cell in row.cells:
                number = extract_ods_number(cell.text)
                if number:
                    return number
    return ""


# --- Utilidades de párrafos ---

def _first_title_text(paragraphs: list[Paragraph]) -> str:
    """Primer texto que no es viñeta ni un encabezado de sección conocido.

    En plantillas sin marcadores, corresponde al título de la actividad.
    """
    markers = (_MARKER_ACTIVITY, _MARKER_ENTREGABLE, _MARKER_REALIZADAS)
    for paragraph in paragraphs:
        if _is_list_paragraph(paragraph):
            continue
        text = paragraph.text.strip()
        if not text:
            continue
        if any(normalize_text(text).startswith(marker) for marker in markers):
            continue
        return text
    return ""


def _is_list_paragraph(paragraph: Paragraph) -> bool:
    """``True`` si el párrafo es una viñeta de lista (por estilo o por numeración).

    Se comprueba el nombre del estilo (tolerante a idioma: "List Paragraph" /
    "Párrafo de lista") y, como señal fuerte, la presencia de ``numPr`` en el XML.
    """
    name = normalize_text(paragraph.style.name or "")
    if "list" in name or "lista" in name:
        return True
    p_pr = paragraph._p.find(qn("w:pPr"))
    return p_pr is not None and p_pr.find(qn("w:numPr")) is not None


def _find_marker(paragraphs: list[Paragraph], marker: str) -> int | None:
    for i, paragraph in enumerate(paragraphs):
        if normalize_text(paragraph.text).startswith(marker):
            return i
    return None


def _text_until_next_marker(paragraphs: list[Paragraph], start: int) -> str:
    """Texto entre el marcador en ``start`` y el siguiente marcador conocido.

    Incluye lo que venga tras los dos puntos del propio marcador (si lo hay) y los
    párrafos siguientes hasta encontrar otro encabezado de sección.
    """
    markers = (_MARKER_ACTIVITY, _MARKER_ENTREGABLE, _MARKER_REALIZADAS)
    pieces: list[str] = []

    first = paragraphs[start].text
    if ":" in first:
        after_colon = first.split(":", 1)[1].strip()
        if after_colon:
            pieces.append(after_colon)

    for paragraph in paragraphs[start + 1 :]:
        normalized = normalize_text(paragraph.text)
        if any(normalized.startswith(m) for m in markers):
            break
        text = paragraph.text.strip()
        if text:
            pieces.append(text)

    return " ".join(pieces).strip()
